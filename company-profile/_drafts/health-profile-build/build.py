# -*- coding: utf-8 -*-
import sys, re, importlib, os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

lang = sys.argv[1]
out = sys.argv[2]
DOC = importlib.import_module("content_" + lang).DOC
RTL = lang == "ar"
ASSETS = os.path.expanduser("~/mnt/syltra smart/company-profile/سيلترا هيلث")
HERO = ASSETS + "/الصور/الهيرو/"
LOCKUP = ASSETS + "/الشعارات/health-lockup-full.png"

SANS = "IBM Plex Sans Arabic" if RTL else "IBM Plex Sans"
MONO = "IBM Plex Mono"
VOID, GRAPHITE, SLATE, CHROME, PLATINUM = "0B0C0E", "17181C", "5B6068", "C7CCD3", "ECEDEF"
GREEN, GREEN_BRIGHT, CARD = "1AA653", "22E06B", "F3F4F6"
AR_RE = re.compile(r"[؀-ۿ]")
PAGE_W = Mm(210); MARG = Mm(20) if RTL else Mm(19); CONTENT_W = PAGE_W - 2 * MARG
SC = 1.0 if RTL else 0.95

def rgb(h): return RGBColor.from_string(h)

def set_run_font(run, family, size, color=VOID, bold=False, rtl=None):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(a), family)
    run.font.size = Pt(size)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs"); rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size * 2)))
    run.font.color.rgb = rgb(color)
    if bold:
        run.font.bold = True
        bCs = OxmlElement("w:bCs"); rPr.append(bCs)
    if rtl is None:
        rtl = RTL and bool(AR_RE.search(run.text))
    if rtl:
        r = OxmlElement("w:rtl"); r.set(qn("w:val"), "1"); rPr.append(r)

def set_par(p, align=None, before=0, after=0, line=1.0, keep_next=False, bidi=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if keep_next: pf.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    if bidi is None: bidi = RTL
    if bidi:
        b = OxmlElement("w:bidi"); b.set(qn("w:val"), "1"); pPr.append(b)
    if align is None: align = "right" if RTL else "left"
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc"); pPr.append(jc)
    jc.set(qn("w:val"), align)

def add_text(container, text, family=None, size=10.5, color=VOID, bold=False, align=None,
             before=0, after=0, line=1.5, keep_next=False, mono=False, rtl=None):
    p = container.add_paragraph()
    fam = MONO if (mono and not AR_RE.search(text)) else (family or SANS)
    latin_only = not AR_RE.search(text)
    bidi = RTL and not (mono and latin_only)
    set_par(p, align=align, before=before, after=after, line=line, keep_next=keep_next, bidi=bidi)
    r = p.add_run(text)
    set_run_font(r, fam, size, color, bold, rtl=rtl)
    return p

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for k, v in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement("w:" + k); e.set(qn("w:w"), str(v)); e.set(qn("w:type"), "dxa"); mar.append(e)
    tcPr.append(mar)

def cell_borders(cell, color="FFFFFF", sz=0, sides=("top", "bottom", "start", "end")):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for s in ("top", "bottom", "start", "end"):
        e = OxmlElement("w:" + s)
        if s in sides and sz:
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz)); e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "nil")
        b.append(e)
    tcPr.append(b)

def keep_table(t):
    rows = t.rows
    for i, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr(); cs = OxmlElement("w:cantSplit"); trPr.append(cs)
        if i < len(rows) - 1:
            for c in row.cells:
                for p in c.paragraphs: p.paragraph_format.keep_with_next = True

def table_setup(t, widths, bidi=None):
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    if bidi is None: bidi = RTL
    if bidi:
        bv = OxmlElement("w:bidiVisual"); tblPr.append(bv)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    # remove default table style borders
    st = tblPr.find(qn("w:tblStyle"))
    if st is not None: tblPr.remove(st)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]
    grid = t._tbl.find(qn("w:tblGrid"))
    for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
        gc.set(qn("w:w"), str(int(widths[i] / 635)))

def clear_cell(cell):
    # remove the default empty paragraph
    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)

def add_image(doc, path, width):
    p = doc.add_paragraph()
    set_par(p, align="center", after=10)
    p.add_run().add_picture(path, width=width)
    return p

def spacer(container, pts):
    p = container.add_paragraph(); set_par(p, after=0, before=0, line=1.0)
    r = p.add_run(""); set_run_font(r, SANS, pts)

# ---------- document ----------
doc = Document()
sec = doc.sections[0]
sec.page_width = PAGE_W; sec.page_height = Mm(297)
sec.left_margin = sec.right_margin = Mm(0); sec.top_margin = sec.bottom_margin = Mm(0)
sec.header_distance = Mm(0); sec.footer_distance = Mm(0)
st = doc.styles["Normal"]; st.font.name = SANS; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:cs"), SANS)

def dark_page(container_doc, build):
    t = container_doc.add_table(rows=1, cols=1)
    table_setup(t, [PAGE_W], bidi=False)
    row = t.rows[0]; row.height = Mm(296); 
    trPr = row._tr.get_or_add_trPr(); h = trPr.find(qn("w:trHeight")); h.set(qn("w:hRule"), "exact")
    c = row.cells[0]; shade(c, VOID); cell_borders(c); cell_margins(c, 1300, 1000, 1130, 1130)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    clear_cell(c); build(c)
    return t

cv = DOC["cover"]
def build_cover(c):
    p = c.add_paragraph(); set_par(p, align=("right" if RTL else "left"), after=0, bidi=False)
    p.add_run().add_picture(LOCKUP, width=Mm(88))
    spacer(c, 150)
    add_text(c, cv["eyebrow"], size=9, color=GREEN_BRIGHT, mono=True, line=1.2, after=18)
    add_text(c, cv["slogan"], size=34, color=PLATINUM, bold=True, line=1.15, after=10)
    add_text(c, cv["sub"], size=15, color=CHROME, line=1.4, after=0)
    spacer(c, 190)
    add_text(c, cv["category"], size=9, color=CHROME, mono=True, line=1.3, after=4)
    add_text(c, cv["endorse"], size=10.5, color=PLATINUM, bold=True, line=1.3)
dark_page(doc, build_cover)

# ---------- body section ----------
body = doc.add_section(WD_SECTION.NEW_PAGE)
body.left_margin = body.right_margin = MARG; body.top_margin = Mm(17); body.bottom_margin = Mm(17)
body.footer_distance = Mm(9)
body.footer.is_linked_to_previous = False
fp = body.footer.paragraphs[0]; set_par(fp, align="center", bidi=False)
r = fp.add_run(DOC["footer"] + "   ·   "); set_run_font(r, MONO, 7.5, SLATE, rtl=False)
r = fp.add_run(); set_run_font(r, MONO, 7.5, SLATE, rtl=False)
for tag, txt in (("begin", None), (None, "PAGE"), ("end", None)):
    if tag:
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), tag)
    else:
        e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = txt
    r._r.append(e)
doc.sections[0].footer.is_linked_to_previous = False

def cards(items, cols):
    rows = (len(items) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    w = [int(CONTENT_W / cols)] * cols
    table_setup(t, w)
    for i, (title, bodytxt) in enumerate(items):
        c = t.rows[i // cols].cells[i % cols]
        shade(c, CARD); cell_borders(c, "FFFFFF", 28, ("top", "bottom", "start", "end")); cell_margins(c, 140, 140, 180, 180)
        clear_cell(c)
        add_text(c, title, size=10.5, bold=True, line=1.3, after=2, keep_next=True)
        add_text(c, bodytxt, size=9*SC, color="2A2D33", line=1.4)
    for i in range(len(items), rows * cols):
        c = t.rows[i // cols].cells[i % cols]; cell_borders(c); clear_cell(c); spacer(c, 4)
    keep_table(t); spacer(doc, 6)

def steps(items):
    t = doc.add_table(rows=len(items), cols=2)
    table_setup(t, [Mm(12), CONTENT_W - Mm(12)])
    for i, (title, bodytxt) in enumerate(items):
        n, c = t.rows[i].cells
        for cc in (n, c):
            cell_borders(cc, PLATINUM, 6, ("bottom",)); cell_margins(cc, 100, 100, 60, 60); clear_cell(cc)
        add_text(n, "%02d" % (i + 1), size=11, color=GREEN, mono=True, line=1.3, before=2, rtl=False)
        add_text(c, title, size=10.5, bold=True, line=1.3, after=1, keep_next=True)
        add_text(c, bodytxt, size=9*SC, color="2A2D33", line=1.4)
    keep_table(t); spacer(doc, 8)

def table(head, rows):
    ncol = len(head)
    t = doc.add_table(rows=len(rows) + 1, cols=ncol)
    if ncol == 3 and head[1] == "":
        w = [Mm(36), Mm(42), CONTENT_W - Mm(78)]
    elif ncol == 3:
        w = [Mm(40), Mm(42), CONTENT_W - Mm(82)]
    else:
        w = [Mm(45), CONTENT_W - Mm(45)]
    table_setup(t, w)
    for j, h in enumerate(head):
        c = t.rows[0].cells[j]; shade(c, GRAPHITE); cell_borders(c); cell_margins(c, 110, 110, 140, 140); clear_cell(c)
        add_text(c, h if h else " ", size=8.5, color=PLATINUM, mono=True, line=1.3)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]; cell_borders(c, PLATINUM, 6, ("bottom",)); cell_margins(c, 120, 120, 140, 140); clear_cell(c)
            mono = (ncol == 3 and head[1] == "" and j == 1)
            add_text(c, v, size=(9.5 if j else 10)*SC, bold=(j == 0), color=(GREEN if mono else ("2A2D33" if j else VOID)), mono=mono, line=1.4)
    keep_table(t); spacer(doc, 8)

def box(title, items):
    t = doc.add_table(rows=1, cols=1); table_setup(t, [CONTENT_W])
    c = t.rows[0].cells[0]; shade(c, GRAPHITE); cell_borders(c); cell_margins(c, 260, 260, 300, 300); clear_cell(c)
    add_text(c, title, size=9, color=GREEN_BRIGHT, mono=True, line=1.3, after=8)
    for it in items:
        add_text(c, it, size=11, color=PLATINUM, line=1.5, after=5)
    keep_table(t); spacer(doc, 8)

for s in DOC["sections"]:
    if s.get("image"):
        p = doc.add_paragraph(); set_par(p, keep_next=True, before=(0 if s["num"] == "01" else 18))
        from PIL import Image
        im = Image.open(HERO + s["image"]); W, H = im.size; h = int(W / 2.9)
        crop = im.crop((0, (H - h) // 2, W, (H - h) // 2 + h)); cp = "/tmp/hero_%s" % s["image"]; crop.save(cp, quality=90)
        p.add_run().add_picture(cp, width=CONTENT_W)
        p.paragraph_format.space_after = Pt(10)
    else:
        spacer(doc, 22)
    p = doc.add_paragraph(); set_par(p, line=1.2, after=6, keep_next=True)
    r = p.add_run(s["num"] + "    "); set_run_font(r, MONO, 9, GREEN, rtl=False)
    r = p.add_run(s["eyebrow"]); set_run_font(r, SANS, 9, GREEN, bold=True)
    add_text(doc, s["title"], size=20, bold=True, line=1.2, after=8, keep_next=True)
    add_text(doc, s["lead"], size=11.5*SC, color="2A2D33", line=1.5, after=8, keep_next=True)
    for b in s["blocks"]:
        k = b["type"]
        if k == "p": add_text(doc, b["text"], size=10*SC, line=1.5, after=6)
        elif k == "small": add_text(doc, b["text"], size=8.5, color=SLATE, line=1.45, after=8)
        elif k == "h2": add_text(doc, b["text"], size=14, bold=True, line=1.3, before=8, after=6, keep_next=True)
        elif k == "quote":
            p = add_text(doc, b["text"], size=17, color=GREEN, bold=True, line=1.35, before=10, after=10)
        elif k == "cards": cards(b["items"], b["cols"])
        elif k == "steps": steps(b["items"])
        elif k == "table": table(b["head"], b["rows"])
        elif k == "box": box(b["title"], b["items"])

# ---------- back cover ----------
back = doc.add_section(WD_SECTION.NEW_PAGE)
back.left_margin = back.right_margin = Mm(0); back.top_margin = back.bottom_margin = Mm(0)
back.footer.is_linked_to_previous = False
for p in back.footer.paragraphs: 
    for r in p.runs: r._r.getparent().remove(r._r)
bk = DOC["back"]
def build_back(c):
    spacer(c, 120)
    add_text(c, bk["title"], size=30, color=PLATINUM, bold=True, line=1.2, after=12)
    add_text(c, bk["body"], size=12.5, color=CHROME, line=1.55, after=26)
    add_text(c, bk["cta"], size=16, color=GREEN_BRIGHT, mono=True, line=1.3, after=40)
    for ln in bk["lines"]:
        add_text(c, ln, size=10, color=CHROME, line=1.5, after=2)
    spacer(c, 150)
    p = c.add_paragraph(); set_par(p, align=("right" if RTL else "left"), after=14, bidi=False)
    p.add_run().add_picture(LOCKUP, width=Mm(60))
    add_text(c, bk["trust"], size=8.5, color=SLATE, line=1.45)
dark_page(doc, build_back)

doc.save(out)
print("saved", out)
